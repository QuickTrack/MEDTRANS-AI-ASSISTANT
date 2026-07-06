"""
Export service: generates professional Word and PDF documents plus
TXT / JSON / CSV exports. Uses python-docx and reportlab.
"""

from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path
from typing import Optional

from config import AppConfig, CONFIG_PATH
from models import Draft, Task, ExportRepository
from services.logger import get_logger

log = get_logger("export")


class ExportService:
    def __init__(self, cfg: Optional[AppConfig] = None) -> None:
        self.cfg = cfg or AppConfig.load(CONFIG_PATH)

    # ---------- helpers ----------
    def _meta(self, task: Optional[Task], draft: Draft) -> dict:
        return {
            "hospital": self.cfg.export.hospital_name or "—",
            "project": (task.project_name if task else None) or "—",
            "clip": (task.clip_number if task else None) or "—",
            "duration": f"{(task.duration_seconds or 0):.1f}s" if task else "—",
            "date": datetime.now().strftime("%Y-%m-%d"),
            "content": draft.content,
        }

    # ---------- Word ----------
    def to_docx(self, task: Optional[Task], draft: Draft, dst: str | Path) -> Path:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        dst = Path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        m = self._meta(task, draft)
        doc = Document()

        # Logo
        logo = self.cfg.export.logo_path
        if logo and Path(logo).exists():
            doc.add_picture(logo, width=Inches(1.5))

        title = doc.add_heading(m["hospital"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for label, key in (("Project", "project"), ("Clip No.", "clip"),
                           ("Duration", "duration"), ("Date", "date")):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            p.add_run(str(m[key]))

        doc.add_heading("Transcript", level=1)
        doc.add_paragraph(m["content"])

        # Footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(f"{m['hospital']} — Confidential").italic = True
        doc.save(str(dst))
        self._record(task, "docx", dst)
        return dst

    # ---------- PDF ----------
    def to_pdf(self, task: Optional[Task], draft: Draft, dst: str | Path) -> Path:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak

        dst = Path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        m = self._meta(task, draft)
        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(
            str(dst), pagesize=A4,
            leftMargin=inch, rightMargin=inch,
            topMargin=inch, bottomMargin=inch,
            title="Medical Transcript",
        )
        story = []
        if self.cfg.export.logo_path and Path(self.cfg.export.logo_path).exists():
            from reportlab.platypus import Image

            story.append(Image(self.cfg.export.logo_path, width=1.2 * inch, height=1.2 * inch))
        story.append(Paragraph(m["hospital"], styles["Title"]))
        story.append(Paragraph(f"Project: {m['project']}", styles["Normal"]))
        story.append(Paragraph(f"Clip No.: {m['clip']}", styles["Normal"]))
        story.append(Paragraph(f"Duration: {m['duration']}", styles["Normal"]))
        story.append(Paragraph(f"Date: {m['date']}", styles["Normal"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Transcript", styles["Heading1"]))
        body = ParagraphStyle("body", parent=styles["Normal"], fontSize=11, leading=15)
        story.append(Paragraph(m["content"].replace("\n", "<br/>"), body))

        def _footer(canvas, d):
            canvas.saveState()
            canvas.setFont("Helvetica", 8)
            canvas.drawCentredString(A4[0] / 2, 0.5 * inch,
                                      f"{m['hospital']} — Confidential   Page {d.page}")
            canvas.restoreState()

        doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
        self._record(task, "pdf", dst)
        return dst

    # ---------- plain text ----------
    def to_txt(self, task: Optional[Task], draft: Draft, dst: str | Path) -> Path:
        dst = Path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        m = self._meta(task, draft)
        content = (
            f"{m['hospital']}\n"
            f"Project: {m['project']}\nClip No.: {m['clip']}\n"
            f"Duration: {m['duration']}\nDate: {m['date']}\n\n"
            f"{m['content']}\n"
        )
        dst.write_text(content, encoding="utf-8")
        self._record(task, "txt", dst)
        return dst

    # ---------- JSON ----------
    def to_json(self, task: Optional[Task], draft: Draft, dst: str | Path) -> Path:
        dst = Path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        data = {
            "hospital": self.cfg.export.hospital_name,
            "project": task.project_name if task else None,
            "clip": task.clip_number if task else None,
            "duration": task.duration_seconds if task else None,
            "date": self._meta(task, draft)["date"],
            "content": draft.content,
            "confidence": draft.confidence,
        }
        dst.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        self._record(task, "json", dst)
        return dst

    # ---------- CSV ----------
    def to_csv(self, task: Optional[Task], draft: Draft, dst: str | Path) -> Path:
        dst = Path(dst)
        dst.parent.mkdir(parents=True, exist_ok=True)
        m = self._meta(task, draft)
        with open(dst, "w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            writer.writerow(["Hospital", "Project", "Clip", "Duration", "Date", "Confidence", "Transcript"])
            writer.writerow([
                m["hospital"], m["project"], m["clip"], m["duration"], m["date"],
                draft.confidence, draft.content,
            ])
        self._record(task, "csv", dst)
        return dst

    def export_all(self, task: Optional[Task], draft: Draft, base_name: str) -> list[Path]:
        out_dir = self.cfg.resolve("exports")
        results = []
        for fmt in self.cfg.export.formats:
            dst = out_dir / f"{base_name}.{fmt}"
            fn = getattr(self, f"to_{fmt}")
            results.append(fn(task, draft, dst))
        return results

    def _record(self, task: Optional[Task], fmt: str, path: Path) -> None:
        try:
            ExportRepository.save(ExportRecord(
                task_id=task.id if task else None, format=fmt, path=str(path),
            ))
        except Exception as exc:  # pragma: no cover
            log.warning("Failed recording export: %s", exc)
